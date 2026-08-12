#!/usr/bin/env python3
"""
Abnahmeskript für das IT-Bestandsaufnahme-Tool.

Erzeugt die Belege selbst, statt sich auf eine Beschreibung zu verlassen.
Wird nach jeder Umsetzung ausgeführt; die vollständige Ausgabe gehört
ungekürzt in den Rechenschaftsbericht.

    ./venv/bin/python abnahme.py

Das Skript arbeitet in einem temporären Datenverzeichnis und fasst
`data/` nicht an. Rückgabewert 0 = alle Prüfungen bestanden, 1 = mindestens
eine Prüfung fehlgeschlagen.
"""

from __future__ import annotations

import io
import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

BASELINE = ROOT / "abnahme_baseline.json"
AUSGABE = ROOT / "exports" / "abnahme"

ergebnisse: list[tuple[str, str, str]] = []  # (id, status, hinweis)


# --------------------------------------------------------------------------
# Hilfen
# --------------------------------------------------------------------------

def block(titel: str) -> None:
    print()
    print("=" * 78)
    print(titel)
    print("=" * 78)


def melde(pruef_id: str, ok: bool | None, hinweis: str) -> None:
    status = "BESTANDEN" if ok else ("NICHT PRÜFBAR" if ok is None else "FEHLGESCHLAGEN")
    ergebnisse.append((pruef_id, status, hinweis))
    print(f"  [{status}] {pruef_id}: {hinweis}")


def schema_felder(typ: str) -> dict:
    from app.services.schema_loader import schema_loader
    schema = schema_loader.get_schema(typ) or {}
    felder = {}
    for abschnitt in schema.get("abschnitte", []):
        for feldef in abschnitt.get("felder", []):
            felder[feldef.get("name")] = feldef
    return felder


# --------------------------------------------------------------------------
# 1 — Testlauf mit Differenz zum letzten Lauf
# --------------------------------------------------------------------------

def pruefung_tests() -> list[str]:
    block("1  TESTLAUF")
    proc = subprocess.run(
        [sys.executable, "-m", "pytest", "-v", "--no-header"],
        cwd=ROOT,
        env={"PYTHONPATH": str(ROOT), "PATH": "/usr/bin:/bin:/usr/local/bin"},
        capture_output=True,
        text=True,
    )
    print(proc.stdout)
    if proc.stderr.strip():
        print(proc.stderr)

    namen = sorted(
        zeile.split(" ")[0].split("::")[-1]
        for zeile in proc.stdout.splitlines()
        if "::" in zeile and (" PASSED" in zeile or " FAILED" in zeile or " ERROR" in zeile)
    )

    alt: list[str] = []
    if BASELINE.exists():
        alt = json.loads(BASELINE.read_text(encoding="utf-8")).get("tests", [])

    neu = sorted(set(namen) - set(alt))
    weg = sorted(set(alt) - set(namen))

    print(f"  Vorher: {len(alt)} Tests   Jetzt: {len(namen)} Tests")
    print(f"  Neu:      {', '.join(neu) if neu else 'keine'}")
    print(f"  Entfernt: {', '.join(weg) if weg else 'keine'}")

    BASELINE.write_text(json.dumps({"tests": namen}, indent=2, ensure_ascii=False), encoding="utf-8")

    if weg:
        melde("T1", False, f"{len(weg)} Test(s) gegenüber dem letzten Lauf entfernt: {', '.join(weg)}")
    else:
        melde("T1", True, "Kein Test wurde entfernt")

    melde("T2", proc.returncode == 0, f"pytest Rückgabewert {proc.returncode}")
    return namen


# --------------------------------------------------------------------------
# 2 — Regelwerks-Linter
# --------------------------------------------------------------------------

def pruefung_linter() -> None:
    block("2  REGELWERKS-LINTER")
    from app.services.rule_engine import rule_engine
    from app.services.schema_loader import schema_loader

    fehler: list[str] = []
    regeln = rule_engine.rules
    print(f"  {len(regeln)} Regeln geladen, {len(schema_loader.schemas)} Schemas geladen")

    # L1: jede Regel verweist auf ein existierendes Feld im richtigen Schema
    for r in regeln:
        rid = r.get("id", "?")
        typ = r.get("gilt_fuer")
        if not typ or typ == "standort":
            continue
        felder = schema_felder(typ)
        for feld in rule_engine._extract_fields(r.get("bedingung", {})):
            if feld == "anzahl_anbindungen":
                continue
            if feld not in felder:
                fehler.append(f"{rid}: Feld '{feld}' existiert nicht in Schema '{typ}'")

    # L2: geprüfter Wert ist im Schema als zulässiger Wert hinterlegt
    def bedingungen(cond):
        if "alle" in cond:
            for s in cond["alle"]:
                yield from bedingungen(s)
        elif "eines" in cond:
            for s in cond["eines"]:
                yield from bedingungen(s)
        elif "feld" in cond:
            yield cond

    for r in regeln:
        rid, typ = r.get("id", "?"), r.get("gilt_fuer")
        if not typ or typ == "standort":
            continue
        felder = schema_felder(typ)
        for cond in bedingungen(r.get("bedingung", {})):
            feldef = felder.get(cond.get("feld"))
            if not feldef or cond.get("operator") != "gleich":
                continue
            zulaessig = [str(w.get("wert")) for w in (feldef.get("werte") or [])]
            if zulaessig and str(cond.get("wert")) not in zulaessig:
                fehler.append(
                    f"{rid}: prüft Wert '{cond.get('wert')}' — Feld '{cond.get('feld')}' "
                    f"kennt nur {zulaessig}"
                )

    # L3: bezeichnung vorhanden, kein Feld 'stufe'
    for r in regeln:
        rid = r.get("id", "?")
        mv = r.get("massnahme_vorschlag") or {}
        if not mv.get("bezeichnung"):
            fehler.append(f"{rid}: massnahme_vorschlag.bezeichnung fehlt")
        if "stufe" in mv or "stufe" in r:
            fehler.append(f"{rid}: verbotenes Feld 'stufe' vorhanden")

    # L4: jedes bewertete Feld hat mindestens eine Regel
    genutzt = {
        f
        for r in regeln
        for f in rule_engine._extract_fields(r.get("bedingung", {}))
    }
    ohne_regel: list[str] = []
    for typ in schema_loader.get_all_types():
        for name, feldef in schema_felder(typ).items():
            if feldef.get("bewertung") and name not in genutzt:
                ohne_regel.append(f"{typ}.{name}")

    # L5: textbaustein ist überall ein Objekt mit feststellung/auswirkung
    alt_format: list[str] = []
    for typ in schema_loader.get_all_types():
        for name, feldef in schema_felder(typ).items():
            for w in feldef.get("werte") or []:
                tb = w.get("textbaustein")
                if tb is None:
                    continue
                if not isinstance(tb, dict):
                    alt_format.append(f"{typ}.{name}={w.get('wert')}")
                elif "feststellung" not in tb:
                    alt_format.append(f"{typ}.{name}={w.get('wert')} (ohne feststellung)")

    for f in fehler:
        print(f"    ! {f}")
    melde("L1", not fehler, f"{len(fehler)} Regelfehler" if fehler else "Alle Regeln verweisen auf gültige Felder und Werte")

    if ohne_regel:
        print(f"    ! bewertete Felder ohne Regel: {', '.join(ohne_regel)}")
    melde("L2", not ohne_regel, f"{len(ohne_regel)} bewertete Felder ohne Regel" if ohne_regel else "Jedes bewertete Feld hat eine Regel")

    if alt_format:
        print(f"    ! textbaustein nicht zweiteilig ({len(alt_format)}): {', '.join(alt_format[:8])}"
              + (" …" if len(alt_format) > 8 else ""))
    melde("L3", not alt_format, f"{len(alt_format)} Textbausteine noch im alten Format" if alt_format else "Alle Textbausteine sind zweiteilig")


# --------------------------------------------------------------------------
# 3 — Referenzauftrag
# --------------------------------------------------------------------------

def referenzauftrag(tmp: Path):
    """Legt einen Auftrag an, der jede Firewall-Regel genau einmal auslöst."""
    from app.models.auftrag import Auftrag
    from app.models.standort import Standort, Internetanbindung
    from app.models.technik import TechnikObjekt
    from app.services.storage import StorageService

    st = StorageService(tmp)

    auftrag = Auftrag(
        id="abnahme-referenz",
        projekt_nummer="ABNAHME-1",
        kunde="Musterfirma GmbH",
        bezeichnung="Referenzauftrag für die Abnahme",
        aktive_bausteine=["firewall", "usv", "serverraum", "netzwerkschrank", "switch", "access_point"],
    )

    sto_a = Standort(
        id="sto-a", auftrag_id=auftrag.id, bezeichnung="Zentrale", ort="München",
        redaktionskonzept_backup_leitung="keine_backup_leitung",
        trassenfuehrung_getrennt="nein",
        usv_fuer_netzwerktechnik="nein",
        anbindungen=[
            Internetanbindung(anbieter="Telekom", art="Glasfaser_FTTH", bandbreite_down_mbit=1000, bandbreite_up_mbit=500),
            Internetanbindung(anbieter="Vodafone", art="Kabel", bandbreite_down_mbit=500, bandbreite_up_mbit=50, ist_backup_leitung="ja"),
            Internetanbindung(anbieter="O2", art="LTE_5G", bandbreite_down_mbit=100, bandbreite_up_mbit=20, ist_backup_leitung="ja"),
        ],
    )
    sto_b = Standort(
        id="sto-b", auftrag_id=auftrag.id, bezeichnung="Aussenstelle", ort="Augsburg",
        anbindungen=[
            Internetanbindung(anbieter="Telekom", art="DSL", bandbreite_down_mbit=100, bandbreite_up_mbit=40),
            Internetanbindung(anbieter="Vodafone", art="DSL", bandbreite_down_mbit=50, bandbreite_up_mbit=10),
        ],
    )
    sto_c = Standort(
        id="sto-c", auftrag_id=auftrag.id, bezeichnung="Filiale", ort="Nürnberg",
        anbindungen=[Internetanbindung(anbieter="Telekom", art="DSL", bandbreite_down_mbit=50, bandbreite_up_mbit=10)],
    )

    schlecht = dict(
        hersteller="Sophos", modell="XGS 2100", aufbau="Einzelgeraet",
        hardware_alter="ueber_5_jahre", wartungsvertrag_vorhanden="nein",
        security_abo_vorhanden="nein", firmware_eol="ja",
        dokumentation_vorhanden="keine", konfigurationssicherung_aktuell="nein",
        konfig_sicherung_automatisch="nein", zugangsschutz_standort="frei_zugaenglich",
        ersatzgeraet_vorhanden="nein", web_protection_aktiv="nein", ips_aktiv="nein",
        mfa_fuer_vpn="nein", vlan_konzept_umgesetzt="nein", exchange_onprem_dahinter="ja",
    )
    gut = dict(
        hersteller="Fortinet", modell="FortiGate 100F", aufbau="Cluster_aktiv_passiv",
        hardware_alter="unter_3_jahre", wartungsvertrag_vorhanden="ja",
        security_abo_vorhanden="ja", firmware_eol="nein",
        dokumentation_vorhanden="vollstaendig", konfigurationssicherung_aktuell="ja",
        konfig_sicherung_automatisch="ja", zugangsschutz_standort="abgeschlossener_raum",
        ersatzgeraet_vorhanden="ja", web_protection_aktiv="ja", ips_aktiv="ja",
        mfa_fuer_vpn="ja", vlan_konzept_umgesetzt="ja", exchange_onprem_dahinter="nein",
    )
    luecken = dict(gut)
    for feld in ("ips_aktiv", "dokumentation_vorhanden", "hardware_alter"):
        luecken.pop(feld, None)

    objekte = [
        TechnikObjekt(id="fw-schlecht", typ="firewall", bezeichnung="Firewall Aussenstelle",
                      auftrag_id=auftrag.id, standort_id=sto_b.id,
                      vertraulichkeit="kundentauglich", daten=schlecht),
        TechnikObjekt(id="fw-gut", typ="firewall", bezeichnung="Firewall Zentrale",
                      auftrag_id=auftrag.id, standort_id=sto_a.id,
                      vertraulichkeit="kundentauglich", daten=gut),
        TechnikObjekt(id="fw-intern", typ="firewall", bezeichnung="Firewall Management",
                      auftrag_id=auftrag.id, standort_id=sto_a.id,
                      vertraulichkeit="intern", daten=luecken),
        TechnikObjekt(id="usv-schlecht", typ="usv", bezeichnung="USV Aussenstelle",
                      auftrag_id=auftrag.id, standort_id=sto_b.id,
                      vertraulichkeit="kundentauglich", daten=dict(
                          hersteller="APC", modell="Smart-UPS 1500", kapazitaet_va=1500,
                          ueberbrueckungszeit_minuten=3, auslastung_prozent=85,
                          batterie_alter="ueber_5_jahre", garantie_geraet_bis="2020-01-01",
                          garantie_batterie_bis="2020-01-01", wartungsvertrag_vorhanden="nein",
                          letzter_batterietest="2020-01-01", abschaltsignal_an_server="nein"
                      )),
        TechnikObjekt(id="srvraum-schlecht", typ="serverraum", bezeichnung="Serverraum Aussenstelle",
                      auftrag_id=auftrag.id, standort_id=sto_b.id,
                      vertraulichkeit="kundentauglich", daten=dict(
                          zugangskontrolle="frei_zugaenglich",
                          zutrittsprotokollierung="nein",
                          zugangsberechtigte_dokumentiert="nein",
                          umweltsensorik="keine",
                          brandmeldeanlage="nein",
                          loeschanlage="keine",
                          stromeinspeisung="eine_einspeisung",
                          notstromaggregat="nein"
                      )),
        TechnikObjekt(id="srvraum-schlecht-2", typ="serverraum", bezeichnung="Serverraum Nebenabteil",
                      auftrag_id=auftrag.id, standort_id=sto_a.id,
                      vertraulichkeit="kundentauglich", daten=dict(
                          zugangskontrolle="schluessel_undokumentiert",
                          umweltsensorik="nur_messung",
                          notstromaggregat="ja",
                          notstromaggregat_letzter_test="2020-01-01"
                      )),
        TechnikObjekt(id="schrank-schlecht", typ="netzwerkschrank", bezeichnung="Netzwerkschrank Aussenstelle",
                      auftrag_id=auftrag.id, standort_id=sto_b.id,
                      vertraulichkeit="kundentauglich", daten=dict(
                          art="offen_im_raum",
                          abschliessbar="nein",
                          belueftung="keine",
                          patchpanel_beschriftung="keine",
                          verkabelungsdokumentation="keine",
                          ausfuehrung_durch="gewachsen_ohne_dokumentation",
                          kabeltyp="cat5",
                          verkabelung_alter="ueber_10_jahre",
                          erweiterungsreserve="keine"
                      )),
        TechnikObjekt(id="sw-schlecht", typ="switch", bezeichnung="Switch Aussenstelle",
                      auftrag_id=auftrag.id, standort_id=sto_b.id,
                      vertraulichkeit="kundentauglich", daten=dict(
                          hersteller="Cisco", management_typ="unmanaged", netztrennung="nein",
                          firmware_aktuell="nein", garantie_bis="2020-01-01",
                          wartungsvertrag_vorhanden="nein", konfigurationssicherung_aktuell="nein",
                          zugangsschutz_management="http_telnet", port_security_aktiv="nein",
                          loop_protection_aktiv="nein"
                      )),
        TechnikObjekt(id="ap-schlecht", typ="access_point", bezeichnung="Access Point Aussenstelle",
                      auftrag_id=auftrag.id, standort_id=sto_b.id,
                      vertraulichkeit="kundentauglich", daten=dict(
                          hersteller="Cisco", wlan_standard="wifi4_oder_aelter", management="standalone",
                          gast_wlan_vorhanden="ja", gast_wlan_isoliert="nein", verschluesselung_wpa3="nein",
                          firmware_aktuell="nein", garantie_bis="2020-01-01", wartungsvertrag_vorhanden="nein"
                      )),
    ]

    st.save_auftrag(auftrag)
    st.save_standort(sto_a)
    st.save_standort(sto_b)
    st.save_standort(sto_c)
    for o in objekte:
        st.save_objekt(o)

    return st, auftrag, [sto_a, sto_b, sto_c], objekte


# --------------------------------------------------------------------------
# 4 — Regelabdeckung
# --------------------------------------------------------------------------

def pruefung_regelabdeckung(st, auftrag, standorte, objekte):
    block("4  REGELABDECKUNG — welche Regel hat ausgelöst")
    from app.services.rule_engine import rule_engine

    findings, offene = rule_engine.evaluate_all(auftrag.id, standorte, objekte, [])
    ausgeloest = {f.quelle for f in findings}

    # Regeln, für deren Baustein der Referenzauftrag kein Objekt enthält,
    # können nicht auslösen. Das ist eine Lücke im Referenzauftrag, kein
    # Regelfehler — deshalb getrennt ausgewiesen.
    vorhandene_typen = {o.typ for o in objekte} | {"standort"}

    stumm_echt: list[str] = []
    stumm_ohne_objekt: list[str] = []

    for r in rule_engine.rules:
        rid = r.get("id", "?")
        typ = r.get("gilt_fuer") or "?"
        if rid in ausgeloest:
            print(f"    ausgelöst      {rid}")
        elif typ in vorhandene_typen:
            print(f"    NICHT AUSGEL.  {rid}   [{typ}]")
            stumm_echt.append(rid)
        else:
            print(f"    kein Objekt    {rid}   [{typ}]")
            stumm_ohne_objekt.append(typ)

    print(f"\n  {len(findings)} Findings, {len(offene)} offene Punkte")
    melde("R1", not stumm_echt,
          "Alle Regeln mit passendem Objekt lösen aus" if not stumm_echt
          else f"{len(stumm_echt)} Regeln lösen trotz vorhandenem Objekt nicht aus: {', '.join(stumm_echt)}")
    melde("R2", len(offene) > 0,
          f"{len(offene)} offene Punkte aus leeren Feldern — leeres Feld erzeugt keinen Mangel")
    if stumm_ohne_objekt:
        melde("R3", None,
              f"{len(stumm_ohne_objekt)} Regeln ungeprüft — Referenzauftrag hat keine Objekte der Typen: "
              f"{', '.join(sorted(set(stumm_ohne_objekt)))}")
    else:
        melde("R3", True, "Der Referenzauftrag deckt alle Bausteine mit Objekten ab")

    for f in findings:
        f.status = "bestaetigt"
    st.save_findings(auftrag.id, findings)
    return findings


# --------------------------------------------------------------------------
# 5 — Exporte und Kennzahlen
# --------------------------------------------------------------------------

def pruefung_exporte(auftrag, standorte, objekte, massnahmen, findings=None):
    block("5  EXPORTE UND KENNZAHLEN")
    from docx import Document
    from app.services.exporter import exporter_service
    from app.services.evaluator import evaluator_service

    AUSGABE.mkdir(parents=True, exist_ok=True)
    zeilen = []
    dokumente = {}

    for stufe in ("intern", "kundentauglich", "anonymisiert"):
        puffer = exporter_service.export_analysebericht_docx(
            auftrag, standorte, objekte, massnahmen, stufe, findings=findings
        )
        pfad = AUSGABE / f"analysebericht_{stufe}.docx"
        pfad.write_bytes(puffer.getvalue())
        doc = Document(str(pfad))
        dokumente[stufe] = doc

        # Objektzeilen aus der Übersichtstabelle zählen
        objektzeilen = 0
        for t in doc.tables:
            kopf = [c.text.strip() for c in t.rows[0].cells]
            if kopf and kopf[0] in ("Bezeichnung", "Objekt", "Objekt / Standort") and "Status" in kopf:
                objektzeilen = len(t.rows) - 1
                break

        _, _, gefiltert, bew = exporter_service._filter_and_evaluate(
            auftrag, standorte, objekte, stufe
        )
        kriterien = sum(len(k.kriterien) for k in bew.kategorien)
        zeilen.append((stufe, len(gefiltert), objektzeilen, kriterien,
                       bew.feldabdeckung_prozent, bew.bausteinabdeckung_prozent,
                       bew.gesamt_prozent, len(doc.inline_shapes)))

    kopf = f"  {'Stufe':<16}{'Objekte':>8}{'Tabelle':>9}{'Krit.':>7}{'Feldabd.':>10}{'Bausteinabd.':>14}{'Score':>8}{'Bilder':>8}"
    print(kopf)
    print("  " + "-" * (len(kopf) - 2))
    for s, n, tz, k, fa, ba, sc, bi in zeilen:
        print(f"  {s:<16}{n:>8}{tz:>9}{k:>7}{fa:>9.1f}%{ba:>13.1f}%{sc:>7.1f}%{bi:>8}")

    intern = next(z for z in zeilen if z[0] == "intern")
    anon = next(z for z in zeilen if z[0] == "anonymisiert")

    melde("E1", all(z[1] > 0 for z in zeilen),
          "Keine Stufe verliert alle Objekte" if all(z[1] > 0 for z in zeilen)
          else "Mindestens eine Stufe enthält 0 Objekte")
    melde("E2", anon[1] == intern[1],
          f"anonymisiert hat {anon[1]} Objekte, intern hat {intern[1]} — müssen gleich sein")
    melde("E3", all(z[1] == z[2] for z in zeilen),
          "Gefilterte Objekte und Tabellenzeilen stimmen überein" if all(z[1] == z[2] for z in zeilen)
          else "Übersichtstabelle zeigt weniger Objekte als der Filter durchlässt — zweite Filterung im Report")
    melde("E4", all(z[7] > 0 for z in zeilen),
          f"Grafik eingebettet (Bilder je Stufe: {[z[7] for z in zeilen]})" if all(z[7] > 0 for z in zeilen)
          else f"Kein Bild im Dokument (Bilder je Stufe: {[z[7] for z in zeilen]})")

    return dokumente


# --------------------------------------------------------------------------
# 6 — Vertraulichkeit inhaltlich
# --------------------------------------------------------------------------

def pruefung_vertraulichkeit(dokumente, objekte):
    block("6  VERTRAULICHKEIT IM TEXT")
    intern_objekte = [o for o in objekte if o.vertraulichkeit == "intern"]

    doc = dokumente["kundentauglich"]
    volltext = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            volltext += "\n" + " | ".join(c.text for c in r.cells)

    lecks = [o.bezeichnung for o in intern_objekte if o.bezeichnung and o.bezeichnung in volltext]
    melde("V1", not lecks,
          "Keine internen Objektbezeichnungen im kundentauglichen Export" if not lecks
          else f"Interne Objekte im kundentauglichen Export: {', '.join(lecks)}")

    anon = dokumente["anonymisiert"]
    anontext = "\n".join(p.text for p in anon.paragraphs)
    for t in anon.tables:
        for r in t.rows:
            anontext += "\n" + " | ".join(c.text for c in r.cells)
    klarnamen = [o.bezeichnung for o in objekte if o.bezeichnung and o.bezeichnung in anontext]
    melde("V2", not klarnamen,
          "Im anonymisierten Export ist keine Originalbezeichnung zu finden" if not klarnamen
          else f"Klarnamen im anonymisierten Export: {', '.join(klarnamen)}")


# --------------------------------------------------------------------------
# 7 — Textbausteine
# --------------------------------------------------------------------------

def pruefung_textbausteine(dokumente):
    block("7  TEXTBAUSTEINE — Feststellung und Auswirkung getrennt")
    from app.services.schema_loader import schema_loader

    paare = []
    for typ in schema_loader.get_all_types():
        for name, feldef in schema_felder(typ).items():
            for w in feldef.get("werte") or []:
                tb = w.get("textbaustein")
                if isinstance(tb, dict) and tb.get("feststellung") and tb.get("auswirkung"):
                    paare.append((tb["feststellung"].strip(), tb["auswirkung"].strip()))

    print(f"  {len(paare)} Textbausteine haben Feststellung UND Auswirkung")
    if not paare:
        melde("B1", None, "Kein Textbaustein hat eine Auswirkung — Trennung nicht prüfbar")
        return

    doc = dokumente["intern"]
    absaetze = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    verklebt = []
    for fest, ausw in paare:
        for a in absaetze:
            if fest[:60] in a and ausw[:60] in a:
                verklebt.append(fest[:50])
                break
    melde("B1", not verklebt,
          "Feststellung und Auswirkung stehen in getrennten Absätzen" if not verklebt
          else f"{len(verklebt)} Textbausteine stehen im selben Absatz")

    # Verklebung über Feldgrenzen: Auswirkung von A + Feststellung von B in einem Absatz
    ueber_grenze = []
    for a in absaetze:
        treffer_fest = sum(1 for f, _ in paare if f[:60] in a)
        treffer_ausw = sum(1 for _, w in paare if w[:60] in a)
        if treffer_fest >= 1 and treffer_ausw >= 1 and (treffer_fest + treffer_ausw) > 1:
            ueber_grenze.append(a[:90])
    melde("B2", not ueber_grenze,
          "Kein Absatz vermischt Texte verschiedener Felder" if not ueber_grenze
          else f"{len(ueber_grenze)} Absätze vermischen Texte verschiedener Felder, z. B.: {ueber_grenze[0]}…")


# --------------------------------------------------------------------------
# 8 — Maßnahmen
# --------------------------------------------------------------------------

def pruefung_massnahmen(st, auftrag, findings):
    block("8  MASSNAHMEN")
    from app.services.rule_engine import rule_engine

    mit_richtwert = [
        r.get("id") for r in rule_engine.rules
        if (r.get("massnahme_vorschlag") or {}).get("kosten_richtwert") is not None
        or (r.get("massnahme_vorschlag") or {}).get("aufwand_richtwert") is not None
    ]
    print(f"  Regeln mit hinterlegtem Richtwert: {len(mit_richtwert)} von {len(rule_engine.rules)}")
    if mit_richtwert:
        melde("M1", True, f"{len(mit_richtwert)} Regeln haben Richtwerte")
    else:
        # Solange die Herkunft der Preise nicht entschieden ist, sind alle
        # Richtwerte bewusst null. Das ist kein Fehler, aber der Übernahmepfad
        # lässt sich im Betrieb dann auch nicht belegen.
        melde("M1", None,
              "Kein Richtwert im Regelwerk gesetzt — derzeit gewollt, Übernahmepfad nur durch Test belegbar")

    quelle = (ROOT / "app" / "web" / "routes_findings.py")
    liest_richtwerte = False
    if quelle.exists():
        text = quelle.read_text(encoding="utf-8")
        liest_richtwerte = "kosten_richtwert" in text and "aufwand_richtwert" in text
    melde("M5", liest_richtwerte,
          "routes_findings.py liest kosten_richtwert und aufwand_richtwert"
          if liest_richtwerte else "Keine lesende Codestelle für die Richtwerte gefunden")

    ohne_bezeichnung = [
        r.get("id") for r in rule_engine.rules
        if not (r.get("massnahme_vorschlag") or {}).get("bezeichnung")
    ]
    melde("M2", not ohne_bezeichnung,
          "Jede Regel hat eine Maßnahmenbezeichnung" if not ohne_bezeichnung
          else f"Ohne Bezeichnung: {', '.join(ohne_bezeichnung)}")

    praefixe = [
        r.get("id") for r in rule_engine.rules
        if ((r.get("massnahme_vorschlag") or {}).get("bezeichnung") or "").startswith("Maßnahme")
    ]
    melde("M3", not praefixe,
          "Keine Bezeichnung trägt ein Präfix" if not praefixe
          else f"Präfix gefunden bei: {', '.join(praefixe)}")

    stufen = [r.get("massnahme_vorschlag", {}) for r in rule_engine.rules]
    melde("M4", not any("stufe" in (mv or {}) for mv in stufen),
          "Kein Regelwerkseintrag enthält noch ein Feld 'stufe'")


# --------------------------------------------------------------------------
# 9 — Kapitelstruktur
# --------------------------------------------------------------------------

def pruefung_kapitel(dokumente):
    block("9  KAPITELSTRUKTUR (intern)")
    doc = dokumente["intern"]
    ueberschriften = [p.text.strip() for p in doc.paragraphs
                      if p.style.name.startswith("Heading") and p.text.strip()]
    for u in ueberschriften:
        print(f"    {u}")

    markdown = [p.text for p in doc.paragraphs if p.text.strip().startswith(("#", "|"))]
    melde("K1", not markdown,
          "Keine Markdown-Reste im Dokument" if not markdown
          else f"{len(markdown)} Absätze enthalten noch Markdown-Zeichen")

    erwartet = ["Feststellung", "Maßnahmen", "Übersicht"]
    fehlend = [e for e in erwartet if not any(e.lower() in u.lower() for u in ueberschriften)]
    melde("K2", not fehlend,
          "Alle Pflichtkapitel vorhanden" if not fehlend
          else f"Fehlende Kapitel: {', '.join(fehlend)}")


# --------------------------------------------------------------------------
# 10 — Inhalt der erzeugten Kapitel
# --------------------------------------------------------------------------

def massnahmen_aus_findings(st, auftrag, findings):
    """Erzeugt aus den bestätigten Findings Maßnahmen, damit Kapitel 6 im
    Abnahmeexport nicht leer bleibt. Stufe wird aus dem Schweregrad abgeleitet."""
    from app.models.massnahme import Massnahme
    from app.services.rule_engine import rule_engine

    stufe_map = {"hoch": 1, "mittel": 2, "niedrig": 3, "empfehlung": 3}
    regeln = {r.get("id"): r for r in rule_engine.rules}
    massnahmen = []

    for i, f in enumerate(findings, start=1):
        regel = regeln.get(f.quelle) or {}
        mv = regel.get("massnahme_vorschlag") or {}
        bezeichnung = mv.get("bezeichnung") or f"Bezeichnung fehlt im Regelwerk (Regel-ID: {f.quelle})"
        massnahmen.append(Massnahme(
            id=f"mn-{i:03d}",
            bezeichnung=bezeichnung,
            beschreibung=mv.get("beschreibung", ""),
            findings=[f.id],
            stufe=stufe_map.get(f.schweregrad, 2),
            prioritaet=f.schweregrad if f.schweregrad in ("hoch", "mittel", "niedrig") else "niedrig",
        ))

    # Rückverweis setzen, damit die Spalte "Maßnahme" in Kapitel 5 gefüllt ist
    # und nicht "Maßnahme ausstehend" zeigt.
    for m in massnahmen:
        for fid in m.findings:
            treffer = next((f for f in findings if f.id == fid), None)
            if treffer is not None:
                treffer.massnahme_id = m.id

    if hasattr(st, "save_massnahmen"):
        st.save_massnahmen(auftrag.id, massnahmen)
    if hasattr(st, "save_findings"):
        st.save_findings(auftrag.id, findings)
    return massnahmen


def _tabelle_mit(doc, *stichworte):
    for t in doc.tables:
        kopf = " | ".join(c.text.strip() for c in t.rows[0].cells)
        if all(s.lower() in kopf.lower() for s in stichworte):
            return t
    return None


def pruefung_inhalt(dokumente, findings, massnahmen):
    block("10  INHALT DER KAPITEL 5 UND 6")
    doc = dokumente["intern"]

    tab_findings = _tabelle_mit(doc, "Feld", "Schweregrad")
    zeilen_f = len(tab_findings.rows) - 1 if tab_findings else 0
    print(f"  Feststellungen: {len(findings)} Findings erzeugt, {zeilen_f} Zeilen im Dokument")
    melde("I1", zeilen_f > 0,
          f"Kapitel 5 enthält {zeilen_f} Zeilen" if zeilen_f > 0
          else "Kapitel 5 ist leer, obwohl Findings vorliegen — die Findings erreichen den Export nicht")

    # Nicht auf "Maßnahme" allein prüfen — die Feststellungstabelle hat eine
    # gleichnamige Spalte und würde hier fälschlich erneut gezählt.
    tab_mn = _tabelle_mit(doc, "Maßnahme", "Priorität", "Investition")
    zeilen_m = len(tab_mn.rows) - 1 if tab_mn else 0
    print(f"  Maßnahmenkatalog: {len(massnahmen)} Maßnahmen übergeben, {zeilen_m} Zeilen im Dokument")
    melde("I2", zeilen_m > 0,
          f"Kapitel 6 enthält {zeilen_m} Zeilen" if zeilen_m > 0
          else "Kapitel 6 ist leer, obwohl Maßnahmen vorliegen")

    if tab_findings and zeilen_f:
        print("  Erste drei Feststellungen:")
        for r in tab_findings.rows[1:4]:
            print("    " + " | ".join(c.text.strip()[:38] for c in r.cells))
    if tab_mn and zeilen_m:
        print("  Erste drei Maßnahmen:")
        for r in tab_mn.rows[1:4]:
            print("    " + " | ".join(c.text.strip()[:38] for c in r.cells))


def pruefung_platzhalter(auftrag, standorte, objekte, massnahmen):
    block("11  PLATZHALTER IM MARKDOWN-EXPORT")
    from app.services.exporter import exporter_service
    try:
        md = exporter_service.export_analysebericht(
            auftrag, standorte, objekte, massnahmen, "kundentauglich"
        )
    except Exception as e:
        melde("P1", None, f"Markdown-Export nicht ausführbar: {e}")
        return

    reste = [z.strip() for z in str(md).splitlines() if "[[" in z and "]]" in z]
    for z in reste:
        print(f"    ! {z}")
    melde("P1", not reste,
          "Keine Platzhalter im Markdown-Export" if not reste
          else f"{len(reste)} Platzhalter stehen im Markdown-Export und damit in der Web-Vorschau")


# --------------------------------------------------------------------------
# 12 — Nachprüfung Fix-Auftrag v6, Stufe 1
# --------------------------------------------------------------------------

def _quelle(pfad: str) -> str:
    p = ROOT / pfad
    return p.read_text(encoding="utf-8") if p.exists() else ""


def pruefung_stufe1():
    """Quelltextprüfungen für die Punkte aus Stufe 1, die sich nicht über den
    Referenzauftrag messen lassen. Bewusst grobe Prüfungen — sie sollen
    verhindern, dass ein bekannter Mangel unbemerkt bestehen bleibt."""
    block("12  NACHPRÜFUNG STUFE 1")

    findings_route = _quelle("app/web/routes_findings.py")
    auftrag_route = _quelle("app/web/routes_auftrag.py")
    findings_tpl = _quelle("app/templates/findings/index.html")
    standort_tpl = _quelle("app/templates/standort/form.html")
    exporter_src = _quelle("app/services/exporter.py")

    # S1 — Sackgasse "uebernommen"
    filter_stellen = []
    for name, text in (("routes_findings.py", findings_route),
                       ("routes_auftrag.py", auftrag_route)):
        for i, z in enumerate(text.splitlines(), 1):
            if 'status != "uebernommen"' in z or "status != 'uebernommen'" in z:
                filter_stellen.append(f"{name}:{i}")
    im_dropdown = "uebernommen" in findings_tpl
    ok_s1 = not filter_stellen and im_dropdown
    melde("S1", ok_s1,
          "Übernommene Findings bleiben sichtbar und sind im Auswahlfeld"
          if ok_s1 else
          f"Filter noch vorhanden: {', '.join(filter_stellen) or 'keiner'}; "
          f"im Auswahlfeld: {'ja' if im_dropdown else 'nein'}")

    # S2 — Vertraulichkeit im Maßnahmenkatalog.
    # Ein bloßes Vorkommen des Parameters genügt nicht: Er muss in einer
    # Bedingung stehen, sonst filtert die Funktion nichts. Genau diese
    # Unterscheidung hat eine frühere, laxere Fassung dieser Prüfung übersehen.
    import ast
    import inspect
    from app.services.exporter import exporter_service

    def _filtert_wirklich(fn) -> tuple[bool, str]:
        try:
            baum = ast.parse(inspect.getsource(fn).lstrip())
        except Exception as e:
            return False, f"Quelltext nicht lesbar ({e})"
        fdef = next((k for k in ast.walk(baum) if isinstance(k, ast.FunctionDef)), None)
        if fdef is None:
            return False, "keine Funktionsdefinition gefunden"
        if not any(a.arg == "ziel_vertraulichkeit" for a in fdef.args.args):
            return False, "kein Vertraulichkeitsparameter"

        relevant = {"ziel_vertraulichkeit", "target_level"}

        # Weg 1: der Wert steht in einer Bedingung, es wird also hier gefiltert.
        in_bedingung = False
        for knoten in ast.walk(fdef):
            teile = []
            if isinstance(knoten, (ast.If, ast.IfExp)):
                teile = [knoten.test]
            elif isinstance(knoten, ast.Compare):
                teile = [knoten]
            elif isinstance(knoten, ast.comprehension):
                teile = list(knoten.ifs)
            for t in teile:
                if {n.id for n in ast.walk(t) if isinstance(n, ast.Name)} & relevant:
                    in_bedingung = True

        # Weg 2: die Daten werden zusammen mit der Stufe an eine Hilfsfunktion
        # gereicht, die filtert. Ein Aufruf, der nur die Stufe umwandelt und das
        # Ergebnis dann liegen lässt, zählt ausdrücklich nicht.
        delegiert = False
        for knoten in ast.walk(fdef):
            if not isinstance(knoten, ast.Call):
                continue
            namen = {n.id for n in ast.walk(knoten) if isinstance(n, ast.Name)}
            if namen & relevant and "massnahmen" in namen:
                delegiert = True

        if not (in_bedingung or delegiert):
            return False, "Stufe wird weder als Bedingung ausgewertet noch zum Filtern weitergereicht"
        return True, ""

    mangel = []
    for fn_name in ("export_massnahmenkatalog_md", "export_massnahmenkatalog_csv"):
        fn = getattr(exporter_service, fn_name, None)
        if fn is None:
            mangel.append(f"{fn_name} fehlt")
            continue
        ok, grund = _filtert_wirklich(fn)
        if not ok:
            mangel.append(f"{fn_name}: {grund}")
    melde("S2", not mangel,
          "Maßnahmenkatalog filtert nach Vertraulichkeit" if not mangel
          else "; ".join(mangel))

    # S3 — automatische Auswertung beim Ansehen.
    # Nur Anzeigerouten (GET) sind gemeint; die ausdrückliche Auswertung über
    # eine POST-Route ist erwünscht und wird nicht bemängelt.
    auto = []
    for name, text in (("routes_findings.py", findings_route),
                       ("routes_auftrag.py", auftrag_route)):
        zeilen = text.splitlines()
        for i, z in enumerate(zeilen, 1):
            if "evaluate_all(" not in z:
                continue
            methode = "?"
            for zurueck in range(i - 1, max(0, i - 60), -1):
                vorher = zeilen[zurueck - 1]
                if "@router." in vorher:
                    methode = vorher.split("@router.")[1].split("(")[0]
                    break
            if methode == "get":
                auto.append(f"{name}:{i}")
    melde("S3", not auto,
          "Keine Regelauswertung in einer Anzeigeroute" if not auto
          else f"Auswertung läuft beim Seitenaufruf: {', '.join(auto)}")

    # S4 — Phantom-Anbindung durch Vorauswahl
    phantom = "not anb and idx > 0" in standort_tpl.replace(" ", " ")
    melde("S4", not phantom,
          "Keine Vorauswahl, die eine leere Anbindung als Eingabe erscheinen lässt"
          if not phantom else
          "standort/form.html wählt bei idx>0 ist_backup_leitung vor — "
          "ein unverändertes Formular erzeugt dadurch eine Anbindung")

    # S5 — Pflichtfeld Begründung
    hat_required = any(
        "begruendung" in z and "required" in z
        for z in findings_tpl.splitlines()
    )
    melde("S5", hat_required,
          "Begründungsfeld ist als Pflichtfeld markiert" if hat_required
          else "Begründungsfeld im Formular ohne required")

    # S6 — Reichweite der Zahlenauswertung
    roh = []
    for p in sorted((ROOT / "app" / "web").glob("*.py")):
        for i, z in enumerate(p.read_text(encoding="utf-8").splitlines(), 1):
            if ".isdigit()" in z:
                roh.append(f"{p.name}:{i}")
    melde("S6", not roh,
          "Keine rohe Zahlenauswertung mehr in den Routen" if not roh
          else f"{len(roh)} Stellen mit .isdigit(): {', '.join(roh)}")

    # S7 — Vollständigkeit der Anonymisierung
    anon_felder = ("strasse", "plz")
    fehlend = [f for f in anon_felder if f not in exporter_src]
    unmaskiert = []
    for f in anon_felder:
        if f'"{f}"' not in exporter_src and f".{f} =" not in exporter_src:
            unmaskiert.append(f)
    melde("S7", not unmaskiert,
          "Anonymisierung erfasst auch Straße und PLZ" if not unmaskiert
          else f"Im anonymisierten Export bleiben im Klartext: {', '.join(unmaskiert)}")


# --------------------------------------------------------------------------
# Hauptlauf
# --------------------------------------------------------------------------

def main() -> int:
    print("ABNAHME — IT-Bestandsaufnahme-Tool")
    print(f"Projekt: {ROOT}")

    pruefung_tests()
    pruefung_linter()

    tmp = Path(tempfile.mkdtemp(prefix="abnahme-data-"))
    try:
        block("3  REFERENZAUFTRAG")
        st, auftrag, standorte, objekte = referenzauftrag(tmp)
        print(f"  Auftrag '{auftrag.id}' mit {len(standorte)} Standorten und {len(objekte)} Objekten")
        print(f"  Temporäres Datenverzeichnis: {tmp}")

        # Findings und Maßnahmen werden über ein modulweites Storage-Objekt
        # nachgeladen. Jedes Modul, das das tut, wird für die Dauer der Prüfung
        # auf das temporäre Verzeichnis umgebogen — sonst liest der Export aus
        # data/, findet nichts, und die Prüfungen laufen ins Leere.
        import importlib
        umgebogen = []
        for modulname in ("app.services.report_builder", "app.services.exporter"):
            try:
                modul = importlib.import_module(modulname)
            except Exception:
                continue
            if hasattr(modul, "storage"):
                umgebogen.append((modul, modul.storage))
                modul.storage = st
        print(f"  Storage umgebogen in: {', '.join(m.__name__ for m, _ in umgebogen) or 'keinem Modul'}")

        try:
            findings = pruefung_regelabdeckung(st, auftrag, standorte, objekte)
            massnahmen = massnahmen_aus_findings(st, auftrag, findings)
            print(f"  {len(massnahmen)} Maßnahmen aus bestätigten Findings erzeugt")
            dokumente = pruefung_exporte(auftrag, standorte, objekte, massnahmen, findings=findings)
            pruefung_inhalt(dokumente, findings, massnahmen)
            pruefung_vertraulichkeit(dokumente, objekte)
            pruefung_textbausteine(dokumente)
            pruefung_massnahmen(st, auftrag, findings)
            pruefung_kapitel(dokumente)
            pruefung_platzhalter(auftrag, standorte, objekte, massnahmen)
            pruefung_stufe1()
        finally:
            for modul, vorher in umgebogen:
                modul.storage = vorher
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    block("ERGEBNIS")
    breite = max(len(i) for i, _, _ in ergebnisse) + 2
    fehl = 0
    for pid, status, hinweis in ergebnisse:
        print(f"  {pid:<{breite}}{status:<16}{hinweis}")
        if status == "FEHLGESCHLAGEN":
            fehl += 1
    print()
    print(f"  {len(ergebnisse)} Prüfungen, {fehl} fehlgeschlagen")
    print(f"  Exporte: {AUSGABE}")
    print()
    return 1 if fehl else 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except Exception:
        import traceback
        traceback.print_exc()
        print("\n  ABBRUCH — das Abnahmeskript ist selbst gescheitert.")
        print("  Das ist ein Befund: Bitte im Rechenschaftsbericht melden, nicht das Skript anpassen.")
        sys.exit(2)
