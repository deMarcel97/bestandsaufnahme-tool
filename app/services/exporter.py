import csv
import io
import json
from typing import List, Dict, Any, Optional
from app.models.auftrag import Auftrag
from app.models.standort import Standort
from app.models.technik import TechnikObjekt
from app.models.finding import Finding
from app.models.massnahme import Massnahme
from app.models.bewertung import GesamtBewertung
from app.services.report_builder import report_builder
from app.services.progress import progress_service
from app.services.evaluator import evaluator_service

from enum import IntEnum
import copy

class VertraulichkeitsStufe(IntEnum):
    ANONYMISIERT = 1
    KUNDENTAUGLICH = 2
    INTERN = 3

    @classmethod
    def parse(cls, val: str) -> "VertraulichkeitsStufe":
        s = str(val or "").strip().lower()
        if s == "anonymisiert":
            return cls.ANONYMISIERT
        elif s == "intern":
            return cls.INTERN
        return cls.KUNDENTAUGLICH

class ExporterService:
    def _filter_and_evaluate(
        self,
        auftrag: Auftrag,
        standorte: List[Standort],
        objekte: List[TechnikObjekt],
        ziel_vertraulichkeit: str
    ):
        target_level = VertraulichkeitsStufe.parse(ziel_vertraulichkeit)

        if target_level == VertraulichkeitsStufe.ANONYMISIERT:
            filtered_objekte = [copy.deepcopy(o) for o in objekte]
            for idx, o in enumerate(filtered_objekte, start=1):
                o.bezeichnung = f"Gerät {idx} ({o.typ.capitalize()})"
                if o.betreut_durch == "Dritter":
                    o.betreut_durch = "[Dienstleister Anonymisiert]"

            filtered_standorte = [copy.deepcopy(s) for s in standorte]
            for idx, s in enumerate(filtered_standorte, start=1):
                s.bezeichnung = f"Standort {idx}"
                s.strasse = "[ANONYMISIERT]"
                s.plz = "[ANONYMISIERT]"
                s.ort = "[ANONYMISIERT]"

            auftrag_copy = copy.deepcopy(auftrag)
            auftrag_copy.kunde = "[ANONYMISIERT]"
        else:
            filtered_objekte = [
                o for o in objekte
                if VertraulichkeitsStufe.parse(getattr(o, "vertraulichkeit", "kundentauglich")) <= target_level
            ]
            filtered_standorte = [
                s for s in standorte
                if VertraulichkeitsStufe.parse(getattr(s, "vertraulichkeit", "kundentauglich")) <= target_level
            ]
            auftrag_copy = auftrag

        filtered_bewertung = evaluator_service.evaluate_auftrag(auftrag.aktive_bausteine, filtered_objekte, filtered_standorte)
        return auftrag_copy, filtered_standorte, filtered_objekte, filtered_bewertung

    def _filter_findings(
        self,
        auftrag_id: str,
        filtered_standorte: List[Standort],
        filtered_objekte: List[TechnikObjekt],
        target_level: VertraulichkeitsStufe,
        provided_findings: Optional[List[Finding]] = None
    ) -> List[Finding]:
        if provided_findings is not None:
            all_findings = provided_findings
        else:
            st = getattr(self, "storage", None)
            if st and hasattr(st, "list_findings"):
                all_findings = st.list_findings(auftrag_id)
            else:
                from app.services.storage import storage
                all_findings = storage.list_findings(auftrag_id)

        if target_level == VertraulichkeitsStufe.ANONYMISIERT:
            return [copy.deepcopy(f) for f in all_findings]

        valid_obj_ids = {o.id for o in filtered_objekte}
        valid_sto_ids = {s.id for s in filtered_standorte}

        filtered_findings = []
        for f in all_findings:
            if f.objekt_id:
                if f.objekt_id in valid_obj_ids:
                    filtered_findings.append(f)
            elif f.standort_id:
                if f.standort_id in valid_sto_ids:
                    filtered_findings.append(f)
            else:
                filtered_findings.append(f)
        return filtered_findings

    def export_analysebericht(
        self,
        auftrag: Auftrag,
        standorte: List[Standort],
        objekte: List[TechnikObjekt],
        massnahmen: List[Massnahme],
        ziel_vertraulichkeit: str = "kundentauglich",
        findings: Optional[List[Finding]] = None
    ) -> str:
        auftrag_copy, filtered_standorte, filtered_objekte, filtered_bewertung = self._filter_and_evaluate(
            auftrag, standorte, objekte, ziel_vertraulichkeit
        )
        target_level = VertraulichkeitsStufe.parse(ziel_vertraulichkeit)
        filtered_findings = self._filter_findings(auftrag.id, filtered_standorte, filtered_objekte, target_level, findings)

        md_text = report_builder.build_analysebericht(
            auftrag_copy, filtered_standorte, filtered_objekte, massnahmen, filtered_bewertung, filtered_findings, ziel_vertraulichkeit
        )
        return md_text.replace("[[GRAFIK:executive_summary]]", "*(Visualisierung der Executive Summary im Word-Export)*")

    def export_analysebericht_docx(
        self,
        auftrag: Auftrag,
        standorte: List[Standort],
        objekte: List[TechnikObjekt],
        massnahmen: List[Massnahme],
        ziel_vertraulichkeit: str = "kundentauglich",
        findings: Optional[List[Finding]] = None
    ) -> io.BytesIO:
        auftrag_copy, filtered_standorte, filtered_objekte, filtered_bewertung = self._filter_and_evaluate(
            auftrag, standorte, objekte, ziel_vertraulichkeit
        )
        target_level = VertraulichkeitsStufe.parse(ziel_vertraulichkeit)
        filtered_findings = self._filter_findings(auftrag.id, filtered_standorte, filtered_objekte, target_level, findings)

        md_text = report_builder.build_analysebericht(
            auftrag_copy, filtered_standorte, filtered_objekte, massnahmen, filtered_bewertung, filtered_findings, ziel_vertraulichkeit
        )

        if "[[GRAFIK:executive_summary]]" not in md_text:
            raise ValueError("DOCX export failed: Placeholder [[GRAFIK:executive_summary]] not found in report text")

        from docx import Document
        from docx.shared import Inches
        from app.services.chart_generator import chart_generator

        doc = Document()
        doc.add_heading(f"Analysebericht: IT-Bestandsaufnahme", 0)

        lines = md_text.split('\n')
        in_table = False
        table_data = []

        def flush_table():
            nonlocal in_table, table_data
            if not table_data:
                return
            clean_rows = []
            for r in table_data:
                cells = [c.strip() for c in r.strip('|').split('|')]
                if any(set(c).issubset({'-', ':', ' '}) for c in cells):
                    continue
                clean_rows.append(cells)
            
            if clean_rows:
                cols_count = max(len(r) for r in clean_rows)
                t = doc.add_table(rows=len(clean_rows), cols=cols_count)
                t.style = 'Table Grid'
                for row_idx, row_cells in enumerate(clean_rows):
                    for col_idx, cell_value in enumerate(row_cells):
                        if col_idx < cols_count:
                            cell = t.cell(row_idx, col_idx)
                            cell.text = cell_value
                            if row_idx == 0:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.bold = True
            in_table = False
            table_data = []

        for line in lines:
            line_str = line.strip()
            
            if line_str == "[[GRAFIK:executive_summary]]":
                if in_table:
                    flush_table()
                try:
                    chart_stream = chart_generator.generate_executive_summary_chart(filtered_bewertung)
                    doc.add_paragraph("Visuelle Auswertung & Executive Overview:")
                    doc.add_picture(chart_stream, width=Inches(6.0))
                except Exception as e:
                    doc.add_paragraph(f"[Hinweis: Grafik konnte nicht gerendert werden: {e}]")
                continue

            if line_str.startswith('|'):
                in_table = True
                table_data.append(line_str)
                continue
            else:
                if in_table:
                    flush_table()

            if not line_str:
                continue

            if line_str.startswith('# '):
                doc.add_heading(line_str[2:].strip(), level=1)
            elif line_str.startswith('## '):
                heading_text = line_str[3:].strip()
                doc.add_heading(heading_text, level=2)
            elif line_str.startswith('### '):
                doc.add_heading(line_str[4:].strip(), level=3)
            elif line_str.startswith('#### '):
                doc.add_heading(line_str[5:].strip(), level=4)
            elif line_str.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                parts = line_str[2:].split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 != 0:
                        run.bold = True
            else:
                p = doc.add_paragraph()
                parts = line_str.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 != 0:
                        run.bold = True

        if in_table:
            flush_table()

        # Automated Post-Export Verification Rules
        for p in doc.paragraphs:
            if p.text.startswith('#'):
                raise ValueError(f"DOCX export verification failed: Paragraph starts with #: '{p.text}'")

        if len(doc.inline_shapes) == 0:
            raise ValueError("DOCX export verification failed: Document contains no embedded charts/images")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def export_massnahmenkatalog_md(self, massnahmen: List[Massnahme], ziel_vertraulichkeit: str = "kundentauglich") -> str:
        target_level = VertraulichkeitsStufe.parse(ziel_vertraulichkeit)
        filtered_massnahmen = massnahmen if ziel_vertraulichkeit else massnahmen
        lines = []
        lines.append("# Maßnahmenkatalog\n")
        
        for stufe in [1, 2, 3]:
            st_m = [m for m in filtered_massnahmen if m.stufe == stufe]
            lines.append(f"## Umsetzungsstufe {stufe}")
            if not st_m:
                lines.append("Keine Maßnahmen zugeordnet.\n")
                continue

            inv = sum(m.investitionskosten for m in st_m)
            mon = sum(m.monatliche_kosten for m in st_m)
            zeit = sum(m.zeitaufwand for m in st_m)

            lines.append("| Bezeichnung | Beschreibung | Priorität | Investition (€) | Monatlich (€) | Zeitaufwand | Status |")
            lines.append("| --- | --- | --- | --- | --- | --- | --- |")
            for m in st_m:
                lines.append(f"| {m.bezeichnung} | {m.beschreibung} | {m.prioritaet} | {m.investitionskosten:.2f} | {m.monatliche_kosten:.2f} | {m.zeitaufwand} {m.zeitaufwand_einheit} | {m.status} |")
            lines.append(f"| **Summe Stufe {stufe}** | | | **{inv:.2f}** | **{mon:.2f}** | **{zeit} Aufwandseinheiten** | |\n")

        return "\n".join(lines)

    def export_massnahmenkatalog_csv(self, massnahmen: List[Massnahme], ziel_vertraulichkeit: str = "kundentauglich") -> str:
        target_level = VertraulichkeitsStufe.parse(ziel_vertraulichkeit)
        filtered_massnahmen = massnahmen if ziel_vertraulichkeit else massnahmen
        output = io.StringIO()
        writer = csv.writer(output, delimiter=";")
        writer.writerow(["Stufe", "Bezeichnung", "Beschreibung", "Prioritaet", "Investitionskosten_EUR", "Monatliche_Kosten_EUR", "Zeitaufwand", "Einheit", "Status"])
        for m in filtered_massnahmen:
            writer.writerow([m.stufe, m.bezeichnung, m.beschreibung, m.prioritaet, f"{m.investitionskosten:.2f}", f"{m.monatliche_kosten:.2f}", m.zeitaufwand, m.zeitaufwand_einheit, m.status])
        return output.getvalue()

    def export_managementsummary(
        self,
        auftrag: Auftrag,
        standorte: List[Standort],
        objekte: List[TechnikObjekt],
        findings: List[Finding],
        massnahmen: List[Massnahme],
        bewertung: Optional[GesamtBewertung] = None,
        ziel_vertraulichkeit: str = "kundentauglich"
    ) -> str:
        auftrag_copy, filtered_standorte, filtered_objekte, filtered_bewertung = self._filter_and_evaluate(
            auftrag, standorte, objekte, ziel_vertraulichkeit
        )
        target_level = VertraulichkeitsStufe.parse(ziel_vertraulichkeit)
        filtered_findings = self._filter_findings(auftrag.id, filtered_standorte, filtered_objekte, target_level, findings)

        lines = []
        kunde_display = "[ANONYMISIERT]" if ziel_vertraulichkeit == "anonymisiert" else auftrag_copy.kunde
        lines.append("# Management Summary")
        lines.append(f"**Kunde:** {kunde_display or 'k.A.'}")
        lines.append(f"**Projekt:** {auftrag_copy.bezeichnung} ({auftrag_copy.projekt_nummer})\n")

        lines.append("## Executive Key Metrics")
        lines.append(f"- **Gesamteinstufung:** {filtered_bewertung.gesamt_stufe_bezeichnung} ({filtered_bewertung.gesamt_prozent} %)")
        lines.append(f"- **Erfassungsgrad:** {filtered_bewertung.erfassungsgrad_prozent} % ({filtered_bewertung.erfassungsgrad_bewertet_anzahl} von {filtered_bewertung.erfassungsgrad_gesamt_anzahl} Kriterien)")
        lines.append(f"- **Standorte Erfasst:** {len(filtered_standorte)}")
        lines.append(f"- **Identifizierte Findings:** {len(filtered_findings)}\n")

        lines.append("## Findings nach Schweregrad")
        counts = {"hoch": 0, "mittel": 0, "niedrig": 0, "empfehlung": 0}
        for f in filtered_findings:
            if f.schweregrad in counts:
                counts[f.schweregrad] += 1
        lines.append(f"- **Hoch:** {counts['hoch']}")
        lines.append(f"- **Mittel:** {counts['mittel']}")
        lines.append(f"- **Niedrig:** {counts['niedrig']}")
        lines.append(f"- **Empfehlung:** {counts['empfehlung']}\n")

        lines.append("## Maßnahmen und Budgetaufwand nach Stufen")
        valid_finding_ids = {f.id for f in filtered_findings}
        filtered_massnahmen = [m for m in massnahmen if not m.findings or any(fid in valid_finding_ids for fid in m.findings)]

        for stufe in [1, 2, 3]:
            st_m = [m for m in filtered_massnahmen if m.stufe == stufe]
            inv = sum(m.investitionskosten for m in st_m)
            mon = sum(m.monatliche_kosten for m in st_m)
            lines.append(f"- **Stufe {stufe} ({len(st_m)} Maßnahmen):** Investition: {inv:.2f} € | Monatlich: {mon:.2f} €")

        return "\n".join(lines)

    def export_offene_punkte_md(
        self,
        auftrag: Auftrag,
        standorte: List[Standort],
        objekte: List[TechnikObjekt],
        rule_open_points: List[Any]
    ) -> str:
        all_op = progress_service.collect_offene_punkte(auftrag, standorte, objekte, rule_open_points)
        lines = []
        lines.append("# Liste Offener Punkte und Fehlender Unterlagen\n")
        lines.append(f"Gesamtanzahl offener Punkte: {len(all_op)}\n")
        lines.append("| ID | Quelle | Beschreibung | Status |")
        lines.append("| --- | --- | --- | --- |")
        for item in all_op:
            lines.append(f"| {item.id} | {item.quelle} | {item.text} | {item.status} |")
        return "\n".join(lines)

    def export_raw_json(
        self,
        auftrag: Auftrag,
        standorte: List[Standort],
        objekte: List[TechnikObjekt],
        findings: List[Finding],
        massnahmen: List[Massnahme],
        ziel_vertraulichkeit: str = "kundentauglich"
    ) -> str:
        auftrag_copy, filtered_standorte, filtered_objekte, filtered_bewertung = self._filter_and_evaluate(
            auftrag, standorte, objekte, ziel_vertraulichkeit
        )
        target_level = VertraulichkeitsStufe.parse(ziel_vertraulichkeit)
        filtered_findings = self._filter_findings(auftrag.id, filtered_standorte, filtered_objekte, target_level, findings)

        valid_finding_ids = {f.id for f in filtered_findings}
        filtered_massnahmen = [m for m in massnahmen if not m.findings or any(fid in valid_finding_ids for fid in m.findings)]

        data = {
            "auftrag": auftrag_copy.model_dump(),
            "standorte": [s.model_dump() for s in filtered_standorte],
            "objekte": [o.model_dump() for o in filtered_objekte],
            "findings": [f.model_dump() for f in filtered_findings],
            "massnahmen": [m.model_dump() for m in filtered_massnahmen]
        }
        return json.dumps(data, indent=2, ensure_ascii=False)

    def _is_accessible(self, item_vertraulichkeit: str, target_vertraulichkeit: str) -> bool:
        return VertraulichkeitsStufe.parse(item_vertraulichkeit) <= VertraulichkeitsStufe.parse(target_vertraulichkeit)

exporter_service = ExporterService()
